import os
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from agent.state import AgentState

def set_cell_background(cell, fill_hex):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), fill_hex)
    tcPr.append(shd)

def doc_builder_node(state: AgentState) -> AgentState:
    state["logs"].append("Building final DOCX document...")
    
    doc = Document()
    
    # Set Margins
    for section in doc.sections:
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)
        
    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)
    font.color.rgb = RGBColor(0x33, 0x33, 0x33) # Off-black
    
    # Document Title
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_p.paragraph_format.space_before = Pt(36)
    title_p.paragraph_format.space_after = Pt(12)
    
    # Try to extract a clean title from request
    title_run = title_p.add_run(f"Project Document: {state['request'][:50]}")
    title_run.font.name = 'Calibri'
    title_run.font.size = Pt(26)
    title_run.font.bold = True
    title_run.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79) # Deep Navy
    
    subtitle_p = doc.add_paragraph()
    subtitle_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_p.paragraph_format.space_after = Pt(36)
    subtitle_run = subtitle_p.add_run("Generated autonomously by LangGraph Agent")
    subtitle_run.font.size = Pt(13)
    subtitle_run.font.italic = True
    subtitle_run.font.color.rgb = RGBColor(0x7F, 0x7F, 0x7F) # Gray
    
    doc.add_page_break()
    
    # Table of Contents placeholder or intro section
    h = doc.add_heading(level=1)
    run = h.add_run("1. Executive Summary & Assumptions")
    run.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)
    h.paragraph_format.space_before = Pt(12)
    h.paragraph_format.space_after = Pt(6)
    
    p = doc.add_paragraph()
    p.add_run("This document was prepared autonomously in response to the user query: ")
    req_run = p.add_run(f"\"{state['request']}\"")
    req_run.font.italic = True
    
    if state["assumptions"]:
        doc.add_heading("Key Assumptions Made", level=2)
        for assumption in state["assumptions"]:
            ap = doc.add_paragraph(style='List Bullet')
            ap.add_run(assumption)
            
    doc.add_page_break()
    
def create_docx_table(doc, rows_data):
    if not rows_data:
        return
    num_cols = len(rows_data[0])
    table = doc.add_table(rows=len(rows_data), cols=num_cols)
    table.style = 'Light Shading Accent 1'
    
    for r_idx, row_cells in enumerate(rows_data):
        row = table.rows[r_idx]
        for c_idx, cell_value in enumerate(row_cells):
            if c_idx < len(row.cells):
                cell = row.cells[c_idx]
                cell.text = cell_value
                if r_idx == 0:
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.font.bold = True
    # Spacing paragraph after table
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(12)

def parse_and_add_markdown(doc, text):
    lines = text.split('\n')
    in_table = False
    table_rows = []
    
    def add_inline_formatting(paragraph, text_str):
        parts = text_str.split('**')
        for i, part in enumerate(parts):
            is_bold = (i % 2 == 1)
            subparts = part.split('*')
            for j, subpart in enumerate(subparts):
                is_italic = (j % 2 == 1)
                run = paragraph.add_run(subpart)
                if is_bold:
                    run.font.bold = True
                if is_italic:
                    run.font.italic = True

    for line in lines:
        line_stripped = line.strip()
        if not line_stripped:
            if in_table and table_rows:
                create_docx_table(doc, table_rows)
                table_rows = []
                in_table = False
            continue
            
        if line_stripped.startswith('|') and line_stripped.endswith('|'):
            if all(c in '|- :+ \t' for c in line_stripped):
                continue
            in_table = True
            cells = [c.strip() for c in line_stripped.split('|')[1:-1]]
            table_rows.append(cells)
            continue
        else:
            if in_table and table_rows:
                create_docx_table(doc, table_rows)
                table_rows = []
                in_table = False

        if line_stripped.startswith('#'):
            depth = 0
            for char in line_stripped:
                if char == '#':
                    depth += 1
                else:
                    break
            heading_content = line_stripped[depth:].strip()
            level = min(max(depth, 1), 4)
            h = doc.add_heading(level=level)
            run = h.add_run(heading_content)
            run.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)
            h.paragraph_format.space_before = Pt(12)
            h.paragraph_format.space_after = Pt(4)
            continue
            
        if line_stripped.startswith('- ') or line_stripped.startswith('* '):
            list_content = line_stripped[2:].strip()
            p = doc.add_paragraph(style='List Bullet')
            p.paragraph_format.space_after = Pt(3)
            add_inline_formatting(p, list_content)
            continue

        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(8)
        p.paragraph_format.line_spacing = 1.15
        add_inline_formatting(p, line_stripped)

    if in_table and table_rows:
        create_docx_table(doc, table_rows)

def doc_builder_node(state: AgentState) -> AgentState:
    state["logs"].append("Building final DOCX document...")
    
    doc = Document()
    
    # Set Margins
    for section in doc.sections:
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)
        
    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)
    font.color.rgb = RGBColor(0x33, 0x33, 0x33) # Off-black
    
    # Document Title
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_p.paragraph_format.space_before = Pt(36)
    title_p.paragraph_format.space_after = Pt(12)
    
    # Try to extract a clean title from request
    title_run = title_p.add_run(f"Project Document: {state['request'][:50]}")
    title_run.font.name = 'Calibri'
    title_run.font.size = Pt(26)
    title_run.font.bold = True
    title_run.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79) # Deep Navy
    
    subtitle_p = doc.add_paragraph()
    subtitle_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_p.paragraph_format.space_after = Pt(36)
    subtitle_run = subtitle_p.add_run("Generated autonomously by LangGraph Agent")
    subtitle_run.font.size = Pt(13)
    subtitle_run.font.italic = True
    subtitle_run.font.color.rgb = RGBColor(0x7F, 0x7F, 0x7F) # Gray
    
    doc.add_page_break()
    
    # Table of Contents placeholder or intro section
    h = doc.add_heading(level=1)
    run = h.add_run("1. Executive Summary & Assumptions")
    run.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)
    h.paragraph_format.space_before = Pt(12)
    h.paragraph_format.space_after = Pt(6)
    
    p = doc.add_paragraph()
    p.add_run("This document was prepared autonomously in response to the user query: ")
    req_run = p.add_run(f"\"{state['request']}\"")
    req_run.font.italic = True
    
    if state["assumptions"]:
        doc.add_heading("Key Assumptions Made", level=2)
        for assumption in state["assumptions"]:
            ap = doc.add_paragraph(style='List Bullet')
            ap.add_run(assumption)
            
    doc.add_page_break()
    
    # Write Sections
    for idx, sec in enumerate(state["sections"]):
        heading_text = sec["heading"]
        body_text = sec["body"]
        
        h = doc.add_heading(level=1)
        run = h.add_run(f"{idx + 2}. {heading_text}")
        run.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)
        h.paragraph_format.space_before = Pt(18)
        h.paragraph_format.space_after = Pt(6)
        
        # Parse and add markdown body text natively
        parse_and_add_markdown(doc, body_text)
            
    # Save the file
    os.makedirs("output", exist_ok=True)
    filename = f"output/document_{os.urandom(4).hex()}.docx"
    doc.save(filename)
    
    state["docx_path"] = filename
    state["logs"].append(f"Successfully saved Word Document to {filename}")
    
    return state

